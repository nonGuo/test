"""
输入文档解析器
解析 RS、TS、Mapping 文档
"""
from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, Tuple
import os
import re


class BaseParser(ABC):
    """文档解析器基类"""

    @abstractmethod
    def parse(self, file_path: str) -> Any:
        pass


class RSParser(BaseParser):
    """
    RS 设计文档解析器
    提取业务背景、测试要点

    支持的文档格式：
    - Word (.docx)

    文档结构假设：
    - 使用标题样式（Heading 1/2/3）划分章节
    - 测试要点章节可能命名为：测试要点、测试设计点、测试关注点等
    - 测试要点内容通常在 List Paragraph 或 Normal 样式中
    """

    # 测试要点章节的可能名称（用于模糊匹配）
    TEST_POINT_KEYWORDS = [
        '测试要点', '测试设计点', '测试关注点', '测试重点',
        '测试项', '测试条目', '测试内容', '测试场景',
        '验证要点', '验证点', '检查点', '检查要点'
    ]

    # 业务背景章节的可能名称
    BACKGROUND_KEYWORDS = [
        '背景', '业务背景', '项目背景', '需求背景',
        '概述', '项目概述', '需求概述', '简介'
    ]

    # 验收标准章节的可能名称
    ACCEPTANCE_KEYWORDS = [
        '验收标准', '验收条件', '验收准则', '交付标准',
        '交付条件', '完成标准', '完成条件'
    ]

    def __init__(self, debug: bool = False):
        """
        初始化解析器

        Args:
            debug: 是否输出调试信息
        """
        self.debug = debug

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        解析 RS 文档

        Args:
            file_path: RS 文档路径 (.docx)

        Returns:
            {
                "title": str,                    # 文档标题
                "business_background": str,      # 业务背景
                "test_points": List[str],        # 测试要点列表
                "acceptance_criteria": List[str], # 验收标准
                "sections": Dict[str, str],      # 其他章节内容
                "raw_content": str               # 原始文本内容
            }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"RS 文档不存在：{file_path}")

        # 根据文件扩展名选择解析方法
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.docx':
            return self._parse_docx(file_path)
        elif ext == '.doc':
            # 尝试使用 docx 解析（可能需要转换）
            try:
                return self._parse_docx(file_path)
            except Exception:
                raise NotImplementedError("请将 .doc 文件转换为 .docx 格式")
        elif ext in ['.txt', '.md']:
            return self._parse_text(file_path)
        else:
            raise ValueError(f"不支持的文件格式：{ext}")

    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")

        doc = Document(file_path)

        # 解析文档结构
        paragraphs = []
        for p in doc.paragraphs:
            paragraphs.append({
                'style': p.style.name if p.style else 'Normal',
                'text': p.text.strip()
            })

        # 解析表格内容（如果有）
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_data.append(cells)
            tables_content.append(table_data)

        # 提取结构化内容
        result = self._extract_structure(paragraphs, tables_content)

        return result

    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        # 简单的段落分割
        lines = content.split('\n')
        paragraphs = []
        for line in lines:
            line = line.strip()
            if line:
                # 尝试识别标题（以 # 开头或全大写或以：结尾）
                if line.startswith('#') or line.isupper() or line.endswith('：') or line.endswith(':'):
                    style = 'Heading 1'
                else:
                    style = 'Normal'
                paragraphs.append({'style': style, 'text': line})

        return self._extract_structure(paragraphs, [])

    def _extract_structure(self, paragraphs: List[Dict], tables: List[List]) -> Dict[str, Any]:
        """
        从段落列表中提取结构化内容

        简化策略：
        1. 提取文档标题
        2. 尝试提取"测试要点"章节内容
        3. 如果找不到，返回原始内容供 LLM 自行理解

        Args:
            paragraphs: 段落列表，每个元素包含 style 和 text
            tables: 表格内容列表

        Returns:
            结构化的文档内容
        """
        raw_content = '\n'.join(p['text'] for p in paragraphs if p['text'])

        result = {
            'title': '',
            'test_points': [],
            'raw_content': raw_content,
            'extraction_method': 'none'  # 记录提取方式
        }

        # 识别章节
        sections = self._identify_sections(paragraphs)

        if self.debug:
            print("\n识别到的章节:")
            for name, content in sections.items():
                preview = content[:50] + '...' if len(content) > 50 else content
                print(f"  {name}: {preview}")

        # 提取标题（通常是第一个 Title 或 Heading 1）
        for p in paragraphs:
            if p['style'] == 'Title' or (p['style'].startswith('Heading') and p['text']):
                result['title'] = p['text']
                break

        # 尝试提取测试要点章节
        test_point_content = None
        for keyword in self.TEST_POINT_KEYWORDS:
            section_name = self._find_section_by_keyword(sections, keyword)
            if section_name:
                test_point_content = sections.get(section_name, '')
                if test_point_content:
                    result['test_points'] = self._parse_test_points(test_point_content)
                    result['extraction_method'] = 'section_matched'
                    result['matched_section'] = section_name
                    if self.debug:
                        print(f"\n[OK] 匹配到测试要点章节: {section_name}")
                    break

        # 如果没找到测试要点章节，尝试从表格提取
        if not result['test_points'] and tables:
            table_points = self._extract_test_points_from_tables(tables)
            if table_points:
                result['test_points'] = table_points
                result['extraction_method'] = 'table_extracted'
                if self.debug:
                    print(f"\n[OK] 从表格提取到 {len(table_points)} 个测试要点")

        # 记录提取状态
        if not result['test_points']:
            result['extraction_method'] = 'raw_content_only'
            if self.debug:
                print("\n[WARN] 未找到测试要点章节，将返回原始内容供 LLM 理解")

        return result

    def _identify_sections(self, paragraphs: List[Dict]) -> Dict[str, str]:
        """
        识别文档章节

        Args:
            paragraphs: 段落列表

        Returns:
            章节名称 -> 章节内容的映射
        """
        sections = {}
        current_section = None
        current_content = []

        for p in paragraphs:
            style = p['style']
            text = p['text']

            if not text:
                continue

            # 判断是否为标题
            is_heading = (
                style.startswith('Heading') or
                style == 'Title' or
                self._is_heading_text(text)
            )

            if is_heading and style != 'Title':
                # 保存上一个章节
                if current_section and current_content:
                    sections[current_section] = '\n'.join(current_content)

                # 开始新章节
                current_section = text
                current_content = []
            else:
                # 添加到当前章节内容
                if current_section:
                    current_content.append(text)

        # 保存最后一个章节
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content)

        return sections

    def _is_heading_text(self, text: str) -> bool:
        """判断文本是否像标题"""
        # 标题通常较短，不以句号结尾
        if len(text) > 50:
            return False
        if text.endswith('。') or text.endswith('.'):
            return False
        # 标题可能包含特定关键词
        heading_keywords = [
            '背景', '概述', '简介', '目的', '范围', '定义',
            '需求', '设计', '实现', '测试', '验收', '部署',
            '说明', '要求', '规范', '流程', '方案'
        ]
        return any(kw in text for kw in heading_keywords)

    def _find_section_by_keyword(self, sections: Dict[str, str], keyword: str) -> Optional[str]:
        """
        根据关键词查找章节

        Args:
            sections: 章节字典
            keyword: 关键词

        Returns:
            匹配的章节名称，未找到返回 None
        """
        # 精确匹配
        if keyword in sections:
            return keyword

        # 模糊匹配
        for section_name in sections.keys():
            if keyword in section_name or section_name in keyword:
                return section_name

        return None

    def _parse_test_points(self, content: str) -> List[str]:
        """
        解析测试要点内容

        Args:
            content: 测试要点章节的文本内容

        Returns:
            测试要点列表
        """
        if not content:
            return []

        points = []

        # 按行分割
        lines = content.split('\n')

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 去除列表标记
            # 常见格式：1. xxx, - xxx, * xxx, • xxx, 等
            cleaned = self._clean_list_item(line)

            if cleaned:
                points.append(cleaned)

        return points

    def _clean_list_item(self, text: str) -> str:
        """清理列表项文本"""
        # 去除常见的列表标记
        patterns = [
            r'^[\d]+[\.、\)\]]\s*',     # 1. 2、 3) 4]
            r'^[a-zA-Z][\.、\)\]]\s*',  # a. b、 c)
            r'^[-*•·]\s+',               # - * • ·
            r'^[○●◇◆□■]\s*',           # 特殊符号
            r'^\([0-9]+\)\s*',          # (1) (2)
        ]

        cleaned = text
        for pattern in patterns:
            cleaned = re.sub(pattern, '', cleaned)

        return cleaned.strip()

    def _parse_list_content(self, content: str) -> List[str]:
        """解析列表形式的内容"""
        return self._parse_test_points(content)  # 使用相同的解析逻辑

    def _extract_test_points_from_tables(self, tables: List[List]) -> List[str]:
        """从表格中提取测试要点"""
        points = []

        for table in tables:
            if not table or len(table) < 2:
                continue

            # 检查表头是否包含测试相关关键词
            header = table[0] if table else []
            header_text = ' '.join(str(cell) for cell in header)

            test_keywords = ['测试', '验证', '检查', '用例', '场景']
            if any(kw in header_text for kw in test_keywords):
                # 提取测试要点列
                for row in table[1:]:  # 跳过表头
                    for cell in row:
                        cell_text = str(cell).strip()
                        if cell_text and len(cell_text) > 5:  # 过滤太短的内容
                            points.append(cell_text)

        return points

    def to_prompt_content(self, result: Dict[str, Any]) -> str:
        """
        将解析结果转换为适合 AI Prompt 的文本格式

        Args:
            result: parse() 方法返回的结果

        Returns:
            格式化的文本内容
        """
        parts = []

        if result.get('title'):
            parts.append(f"【文档标题】{result['title']}")

        if result.get('test_points'):
            points_str = '\n'.join(f"- {p}" for p in result['test_points'])
            parts.append(f"【测试要点】\n{points_str}")

        return '\n\n'.join(parts)

    def extract_with_llm(self, file_path: str, llm_client: Any) -> Dict[str, Any]:
        """
        使用 LLM 提取测试要点（回退方案）

        Args:
            file_path: RS 文档路径
            llm_client: LLM 客户端

        Returns:
            解析结果
        """
        # 先进行基础解析
        result = self.parse(file_path)

        # 如果已经提取到测试要点，直接返回
        if result.get('test_points'):
            return result

        # 使用 LLM 提取测试要点
        raw_content = result.get('raw_content', '')

        if not raw_content:
            return result

        # 构建 LLM Prompt
        prompt = f"""请从以下 RS 需求文档中提取测试要点。

文档内容：
{raw_content[:3000]}

请提取测试要点，输出 JSON 格式：
{{
  "test_points": ["测试要点1", "测试要点2", ...]
}}

注意：
1. 只提取与测试相关的要点
2. 不要编造不存在的内容
3. 每个测试要点应该简洁明确
"""

        try:
            import json
            import re

            response = llm_client.generate(prompt)

            # 解析 JSON
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                extracted = json.loads(json_match.group())
                result['test_points'] = extracted.get('test_points', [])
                result['extraction_method'] = 'llm_extracted'

                if self.debug:
                    print(f"\n[OK] LLM 提取到 {len(result['test_points'])} 个测试要点")

        except Exception as e:
            if self.debug:
                print(f"\n[WARN] LLM 提取失败: {e}")

        return result


class DWSTableMetadata:
    """
    DWS 分布式数据库表元数据
    
    对象类型说明：
    - TMP: 临时表，中间加工过程
    - FACT: F表，事实表，承载加工好的逻辑
    - INTERFACE: I接口，对外提供数据，多为视图，底层对应F表
    """
    
    def __init__(self):
        # 基本信息
        self.table_name: str = ""           # 表名
        self.table_type: str = ""           # 表类型: TMP/FACT/INTERFACE
        self.schema_name: str = ""          # Schema 名称
        
        # I接口特有属性
        self.underlying_f_table: str = ""   # 底层F表 (I接口对应)
        self.is_view: bool = False          # 是否为视图
        
        # 分布配置 (仅 F表/TMP表 需要，视图无分布配置)
        self.distribution_type: str = ""    # 分布方式: HASH/REPLICATE/RANDOM/NONE
        self.distribution_key: List[str] = []  # 分布键
        
        # 分区配置
        self.partition_type: str = ""       # 分区方式: RANGE/LIST/HASH/NONE
        self.partition_keys: List[str] = [] # 分区键
        self.partition_spec: str = ""       # 分区规格描述
        
        # 存储配置
        self.storage_format: str = ""       # 存储格式: ORC/PARQUET/TEXTFILE
        self.compression: str = ""          # 压缩方式: SNAPPY/GZIP/NONE
        
        # 依赖关系
        self.source_tables: List[str] = []  # 来源表
        self.dim_tables: List[str] = []     # 关联维度表 (I接口特有)
        self.temp_tables: List[str] = []    # 加工过程中的临时表
        
        # 约束
        self.primary_keys: List[str] = []   # 主键
        
        # 其他
        self.etl_job_name: str = ""         # ETL 作业名称
        self.description: str = ""          # 表描述
    
    def to_dict(self) -> Dict[str, Any]:
        """转换为字典"""
        return {
            'table_name': self.table_name,
            'table_type': self.table_type,
            'schema_name': self.schema_name,
            'underlying_f_table': self.underlying_f_table,
            'is_view': self.is_view,
            'distribution_type': self.distribution_type,
            'distribution_key': self.distribution_key,
            'partition_type': self.partition_type,
            'partition_keys': self.partition_keys,
            'partition_spec': self.partition_spec,
            'storage_format': self.storage_format,
            'compression': self.compression,
            'source_tables': self.source_tables,
            'dim_tables': self.dim_tables,
            'temp_tables': self.temp_tables,
            'primary_keys': self.primary_keys,
            'etl_job_name': self.etl_job_name,
            'description': self.description,
        }
    
    @classmethod
    def from_dict(cls, data: Dict[str, Any]) -> 'DWSTableMetadata':
        """从字典创建"""
        obj = cls()
        obj.table_name = data.get('table_name', '')
        obj.table_type = data.get('table_type', '')
        obj.schema_name = data.get('schema_name', '')
        obj.underlying_f_table = data.get('underlying_f_table', '')
        obj.is_view = data.get('is_view', False)
        obj.distribution_type = data.get('distribution_type', '')
        obj.distribution_key = data.get('distribution_key', [])
        obj.partition_type = data.get('partition_type', '')
        obj.partition_keys = data.get('partition_keys', [])
        obj.partition_spec = data.get('partition_spec', '')
        obj.storage_format = data.get('storage_format', '')
        obj.compression = data.get('compression', '')
        obj.source_tables = data.get('source_tables', [])
        obj.dim_tables = data.get('dim_tables', [])
        obj.temp_tables = data.get('temp_tables', [])
        obj.primary_keys = data.get('primary_keys', [])
        obj.etl_job_name = data.get('etl_job_name', '')
        obj.description = data.get('description', '')
        return obj
    
    def get_test_target_table(self) -> str:
        """
        获取测试目标表
        
        测试策略：
        - I接口: 测试对外数据准确性，返回 I 接口名
        - F表: 测试分布/分区配置，返回 F 表名
        - TMP表: 中间过程验证，返回 TMP 表名
        """
        return self.table_name
    
    def get_distribution_check_table(self) -> str:
        """
        获取分布方式检查的表
        
        分布配置检查针对 F 表：
        - I接口: 返回底层 F 表
        - F表: 返回自身
        - TMP表: 返回自身
        """
        if self.table_type == 'INTERFACE' and self.underlying_f_table:
            return self.underlying_f_table
        return self.table_name


class TSParser(BaseParser):
    """
    TS 表模型设计文档解析器
    
    解析 DWS 数据库表对象的元数据信息，包括：
    - 表基本信息（类型、名称、Schema）
    - 分布配置（分布方式、分布键）
    - 分区配置（分区方式、分区键）
    - 依赖关系（来源表、临时表、维度表）
    
    使用 LLM 从 TS 文档中提取结构化信息
    """
    
    # JSON Schema 定义，用于指导 LLM 输出
    OUTPUT_SCHEMA = """
{
  "tables": [
    {
      "table_name": "表名",
      "table_type": "TMP/FACT/INTERFACE",
      "schema_name": "Schema名称",
      "underlying_f_table": "底层F表名(仅I接口需要)",
      "is_view": true/false,
      "distribution_type": "HASH/REPLICATE/RANDOM/NONE",
      "distribution_key": ["分布键1", "分布键2"],
      "partition_type": "RANGE/LIST/HASH/NONE",
      "partition_keys": ["分区键1", "分区键2"],
      "partition_spec": "分区规格描述",
      "storage_format": "ORC/PARQUET/TEXTFILE",
      "compression": "SNAPPY/GZIP/NONE",
      "source_tables": ["来源表1", "来源表2"],
      "dim_tables": ["维度表1"],
      "temp_tables": ["临时表1", "临时表2"],
      "primary_keys": ["主键字段"],
      "etl_job_name": "ETL作业名",
      "description": "表描述"
    }
  ]
}
"""
    
    def __init__(self, debug: bool = False):
        self.debug = debug
        self.tables: List[DWSTableMetadata] = []
    
    def parse(self, file_path: str, llm_client: Any = None) -> Dict[str, Any]:
        """
        解析 TS 文档
        
        Args:
            file_path: TS 文档路径 (.docx/.doc/.txt)
            llm_client: LLM 客户端（必须提供，用于提取结构化信息）
        
        Returns:
            {
                "tables": List[DWSTableMetadata],
                "interface_table": DWSTableMetadata,  # 主 I 接口
                "fact_table": DWSTableMetadata,        # 主 F 表
                "temp_tables": List[DWSTableMetadata], # 临时表列表
                "raw_content": str
            }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"TS 文档不存在：{file_path}")
        
        # 读取文档内容
        raw_content = self._read_document(file_path)
        
        if self.debug:
            print(f"[INFO] 文档内容长度: {len(raw_content)} 字符")
        
        # 必须提供 LLM 客户端
        if llm_client is None:
            if self.debug:
                print("[WARN] 未提供 LLM 客户端，返回原始内容")
            return {
                "tables": [],
                "interface_table": None,
                "fact_table": None,
                "temp_tables": [],
                "raw_content": raw_content
            }
        
        # 使用 LLM 提取结构化信息
        tables = self._extract_with_llm(raw_content, llm_client)
        
        # 分类整理
        result = self._classify_tables(tables)
        result['raw_content'] = raw_content
        
        return result
    
    def _read_document(self, file_path: str) -> str:
        """读取文档内容"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self._read_docx(file_path)
        elif ext in ['.txt', '.md']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            # 尝试作为 docx 读取
            try:
                return self._read_docx(file_path)
            except Exception:
                raise ValueError(f"不支持的文件格式：{ext}")
    
    def _read_docx(self, file_path: str) -> str:
        """读取 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("需要安装 python-docx：pip install python-docx")
        
        doc = Document(file_path)
        
        # 提取段落文本
        paragraphs = []
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                paragraphs.append(text)
        
        # 提取表格文本
        tables_text = []
        for i, table in enumerate(doc.tables):
            table_lines = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            if table_lines:
                tables_text.append(f"[表格 {i+1}]\n" + "\n".join(table_lines))
        
        content = "\n\n".join(paragraphs)
        if tables_text:
            content += "\n\n" + "\n\n".join(tables_text)
        
        return content
    
    def _extract_with_llm(self, content: str, llm_client: Any) -> List[DWSTableMetadata]:
        """使用 LLM 提取表元数据"""
        
        # 限制内容长度
        max_length = 6000
        if len(content) > max_length:
            content = content[:max_length] + "\n... (内容已截断)"
        
        prompt = f"""你是一位 DWS 数据仓库专家。请从以下 TS 表模型设计文档中提取表对象的元数据信息。

文档内容：
{content}

请按以下 JSON 格式输出：

{self.OUTPUT_SCHEMA}

提取规则：
1. table_type 判断规则：
   - 表名以 _I 结尾 → INTERFACE (I接口)
   - 表名以 _F 结尾 → FACT (事实表)
   - 表名以 TMP 或 tmp 开头/包含 → TMP (临时表)

2. I接口特有：
   - 如果是视图，is_view = true
   - underlying_f_table 填写对应的 F 表名

3. 分布配置：
   - HASH: 哈希分布，需要填写 distribution_key
   - REPLICATE: 复制分布，无分布键
   - RANDOM: 随机分布
   - 视图无分布配置

4. 分区配置：
   - RANGE: 范围分区（如按日期）
   - LIST: 列表分区
   - HASH: 哈希分区
   - partition_spec 描述分区细节

5. 依赖关系：
   - source_tables: 数据来源表
   - dim_tables: 关联的维度表
   - temp_tables: 加工过程中的临时表

6. 只输出 JSON，不要包含其他解释。
"""
        
        try:
            import json
            import re
            
            response = llm_client.generate(prompt)
            
            if self.debug:
                print(f"[INFO] LLM 响应长度: {len(response)} 字符")
            
            # 解析 JSON
            json_match = re.search(r'\{[\s\S]*\}', response)
            if json_match:
                data = json.loads(json_match.group())
                tables = []
                
                for table_data in data.get('tables', []):
                    table = DWSTableMetadata.from_dict(table_data)
                    tables.append(table)
                
                if self.debug:
                    print(f"[OK] 提取到 {len(tables)} 个表对象")
                
                return tables
            
        except Exception as e:
            if self.debug:
                print(f"[ERROR] LLM 提取失败: {e}")
        
        return []
    
    def _classify_tables(self, tables: List[DWSTableMetadata]) -> Dict[str, Any]:
        """分类整理表对象"""
        interface_tables = []
        fact_tables = []
        temp_tables = []
        
        for table in tables:
            if table.table_type == 'INTERFACE':
                interface_tables.append(table)
            elif table.table_type == 'FACT':
                fact_tables.append(table)
            elif table.table_type == 'TMP':
                temp_tables.append(table)
        
        # 主 I 接口（通常只有一个）
        main_interface = interface_tables[0] if interface_tables else None
        
        # 主 F 表（与 I 接口关联的）
        main_fact = None
        if main_interface and main_interface.underlying_f_table:
            for f in fact_tables:
                if f.table_name == main_interface.underlying_f_table:
                    main_fact = f
                    break
        if not main_fact:
            main_fact = fact_tables[0] if fact_tables else None
        
        return {
            'tables': tables,
            'interface_table': main_interface,
            'fact_table': main_fact,
            'temp_tables': temp_tables
        }
    
    def to_prompt_content(self, result: Dict[str, Any]) -> str:
        """
        将解析结果转换为适合 AI Prompt 的文本格式
        
        Args:
            result: parse() 方法返回的结果
        
        Returns:
            格式化的文本内容
        """
        parts = []
        
        # I 接口信息
        interface = result.get('interface_table')
        if interface:
            parts.append(f"【I接口】{interface.table_name}")
            if interface.is_view:
                parts.append("  类型: 视图")
            if interface.underlying_f_table:
                parts.append(f"  底层F表: {interface.underlying_f_table}")
            if interface.source_tables:
                parts.append(f"  来源表: {', '.join(interface.source_tables)}")
            if interface.dim_tables:
                parts.append(f"  关联维度: {', '.join(interface.dim_tables)}")
            if interface.primary_keys:
                parts.append(f"  主键: {', '.join(interface.primary_keys)}")
        
        # F 表信息
        fact = result.get('fact_table')
        if fact:
            parts.append(f"\n【F事实表】{fact.table_name}")
            if fact.distribution_type:
                parts.append(f"  分布方式: {fact.distribution_type}")
            if fact.distribution_key:
                parts.append(f"  分布键: {', '.join(fact.distribution_key)}")
            if fact.partition_type:
                parts.append(f"  分区方式: {fact.partition_type}")
            if fact.partition_keys:
                parts.append(f"  分区键: {', '.join(fact.partition_keys)}")
            if fact.partition_spec:
                parts.append(f"  分区规格: {fact.partition_spec}")
            if fact.storage_format:
                parts.append(f"  存储格式: {fact.storage_format}")
        
        # 临时表
        temp_tables = result.get('temp_tables', [])
        if temp_tables:
            parts.append(f"\n【临时表】")
            for t in temp_tables[:5]:  # 最多显示5个
                parts.append(f"  - {t.table_name}")
                if t.distribution_type:
                    parts.append(f"    分布: {t.distribution_type}")
        
        return '\n'.join(parts)


class MappingParserWrapper(BaseParser):
    """
    Mapping 文档解析器包装器
    使用 mapping_parser.py 中的实现
    """

    def __init__(self, debug: bool = False):
        self.debug = debug

    def parse(self, file_path: str) -> List[Dict[str, Any]]:
        """
        解析 Mapping 文档

        Returns:
            [
                {
                    "source_table": str,
                    "source_field": str,
                    "target_table": str,
                    "target_field": str,
                    "transformation_rule": str,
                    "rule_type": str  # direct/calculation/case/join
                }
            ]
        """
        from .mapping_parser import MappingParser

        parser = MappingParser(debug=self.debug)
        result = parser.parse(file_path)

        # 转换为简化格式
        mappings = []
        for fm in result.get('field_mappings', []):
            mappings.append({
                'source_table': fm.get('source_table', ''),
                'source_field': fm.get('source_field', ''),
                'target_table': fm.get('target_table', ''),
                'target_field': fm.get('target_field', ''),
                'transformation_rule': fm.get('transform_rule', ''),
                'rule_type': self._infer_rule_type(fm.get('mapping_scene', ''))
            })

        return mappings

    def _infer_rule_type(self, mapping_scene: str) -> str:
        """推断规则类型"""
        if not mapping_scene:
            return 'direct'

        scene_map = {
            '直接复制': 'direct',
            '直接映射': 'direct',
            '直取': 'direct',
            '数据加工': 'calculation',
            '计算': 'calculation',
            '聚合': 'aggregation',
            '关联': 'join',
            '赋值': 'constant',
            '常量': 'constant',
        }

        return scene_map.get(mapping_scene, 'direct')


class DocumentParserFactory:
    """文档解析器工厂"""

    @staticmethod
    def get_parser(doc_type: str, debug: bool = False) -> BaseParser:
        """
        获取解析器

        Args:
            doc_type: RS/TS/MAPPING
            debug: 是否输出调试信息

        Returns:
            对应的解析器实例
        """
        parsers = {
            "RS": RSParser(debug=debug),
            "TS": TSParser(),
            "MAPPING": MappingParserWrapper(debug=debug)
        }
        return parsers.get(doc_type.upper())

    @staticmethod
    def parse_all(rs_path: str, ts_path: str, mapping_path: str,
                  debug: bool = False) -> Dict[str, Any]:
        """
        解析所有文档

        Returns:
            包含所有文档内容的字典
        """
        rs_parser = DocumentParserFactory.get_parser("RS", debug=debug)
        ts_parser = DocumentParserFactory.get_parser("TS", debug=debug)
        mapping_parser = DocumentParserFactory.get_parser("MAPPING", debug=debug)

        return {
            "rs": rs_parser.parse(rs_path),
            "ts": ts_parser.parse(ts_path),
            "mapping": mapping_parser.parse(mapping_path)
        }