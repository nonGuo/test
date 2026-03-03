# -*- coding: utf-8 -*-
"""分析 RS 文档结构"""
from docx import Document
import json

doc = Document('RS样例.docx')

result = {
    'paragraphs': [],
    'tables': []
}

# 分析段落
for i, p in enumerate(doc.paragraphs):
    result['paragraphs'].append({
        'index': i,
        'style': p.style.name if p.style else 'None',
        'text': p.text
    })

# 分析表格
for ti, table in enumerate(doc.tables):
    table_data = {
        'index': ti,
        'rows': len(table.rows),
        'cols': len(table.columns),
        'data': []
    }
    for ri, row in enumerate(table.rows):
        cells = [cell.text for cell in row.cells]
        table_data['data'].append(cells)
    result['tables'].append(table_data)

# 输出 JSON
with open('rs_structure.json', 'w', encoding='utf-8') as f:
    json.dump(result, f, ensure_ascii=False, indent=2)

print('Done! Saved to rs_structure.json')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')